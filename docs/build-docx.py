# -*- coding: utf-8 -*-
"""Genera docs/Valeria-Manual-Casos-de-Uso.docx A PARTIR del HTML del manual.

    pip install python-docx lxml
    python3 docs/build-docx.py

Antes este script llevaba el texto del manual DUPLICADO dentro, con un aviso en
la cabecera que decía «todo cambio hay que aplicarlo en los dos sitios o el Word
y el PDF se quedan desincronizados sin que nada avise». Es exactamente lo que
pasó: el HTML llegó a la v12 y el Word seguía describiendo la v10.3, con la
mascota antigua y siete niveles.

Ahora hay una sola fuente, docs/manual-casos-de-uso.html, y de ella salen los
tres formatos: HTML, PDF (docs/build-pdf.js) y este Word. Lo que no se puede
representar en Word —los círculos de la portada, los degradados— se sustituye
por su equivalente sobrio; el texto es el mismo, siempre.
"""
import os
import re

import lxml.html
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, 'docs')
SRC = os.path.join(DOCS, 'manual-casos-de-uso.html')
OUT = os.path.join(DOCS, 'Valeria-Manual-Casos-de-Uso.docx')

PRIMARY = RGBColor(0x00, 0xA3, 0x9E)
VIOLET_DARK = RGBColor(0x6D, 0x3F, 0xC4)
INK = RGBColor(0x1F, 0x29, 0x37)
INK2 = RGBColor(0x4B, 0x55, 0x63)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WARN_INK = RGBColor(0x92, 0x5E, 0x0B)
OK_INK = RGBColor(0x0F, 0x76, 0x4E)

FILL_HEAD = 'E6F9F8'
FILL_ZEBRA = 'F7FBFB'
CALLOUT_FILL = {'': 'F0FDF9', 'warn': 'FFFBEB', 'ok': 'EAFAF2', 'violet': 'F5F0FF'}
CALLOUT_INK = {'': PRIMARY, 'warn': WARN_INK, 'ok': OK_INK, 'violet': VIOLET_DARK}

# El ancho útil de la caja de texto: A4 menos los dos márgenes de 2 cm.
CONTENT_CM = 17.0


# --------------------------------------------------------------------------- util
def shade(cell_or_par, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor)
    target = cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, '_tc') \
        else cell_or_par._p.get_or_add_pPr()
    target.append(el)


def no_borders(table):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    table._tbl.tblPr.append(borders)


def classes(el):
    return (el.get('class') or '').split()


def clean(text):
    """Colapsa el espacio en blanco del HTML como haría el navegador."""
    return re.sub(r'\s+', ' ', text or '')


class Builder:
    def __init__(self, doc):
        self.doc = doc
        self.figura = 0

    # ---------------------------------------------------------------- runs
    def runs(self, par, el, bold=False, italic=False, color=None, size=None, mono=False):
        """Vuelca el contenido en línea de `el` (texto + hijos + colas) en `par`."""
        def emit(txt, **kw):
            if not txt:
                return
            r = par.add_run(txt)
            r.bold = kw.get('bold', bold)
            r.italic = kw.get('italic', italic)
            c = kw.get('color', color)
            if c is not None:
                r.font.color.rgb = c
            s = kw.get('size', size)
            if s is not None:
                r.font.size = Pt(s)
            if kw.get('mono', mono):
                r.font.name = 'Consolas'

        emit(clean(el.text))
        for child in el:
            tag = child.tag
            cls = classes(child)
            if tag == 'br':
                par.add_run().add_break()
            elif tag in ('strong', 'b'):
                self.runs(par, child, bold=True, italic=italic, color=color, size=size)
            elif tag in ('em', 'i'):
                self.runs(par, child, bold=bold, italic=True, color=color, size=size)
            elif tag == 'code':
                self.runs(par, child, bold=bold, italic=italic, color=color, size=size, mono=True)
            elif tag == 'span' and 'bdot' in cls:
                # El punto de color del bloque: en Word, un cuadrado del acento.
                m = re.search(r'background:(#[0-9a-fA-F]{6})', child.get('style') or '')
                r = par.add_run('■ ')
                if m:
                    r.font.color.rgb = RGBColor.from_string(m.group(1)[1:].upper())
            elif tag == 'span' and 'fig' in cls:
                self.figura += 1
                emit('Fig. %d' % self.figura, bold=True, color=PRIMARY)
            elif tag == 'span' and ('small' in cls or 'muted' in cls):
                self.runs(par, child, bold=bold, italic=italic,
                          color=MUTED if 'muted' in cls else color, size=8.5)
            elif tag == 'span' and 'stars' in cls:
                self.runs(par, child, bold=True, color=RGBColor(0xEA, 0xB3, 0x08))
            elif tag == 'img':
                # El retrato de Lúa del capítulo 1 va dentro de una tabla de
                # maquetación: si se descarta aquí, el Word se queda sin la
                # mascota justo en la página que la presenta.
                src = os.path.join(DOCS, child.get('src') or '')
                if os.path.exists(src):
                    par.add_run().add_picture(src, width=Cm(2.8))
            else:
                self.runs(par, child, bold=bold, italic=italic, color=color, size=size)
            emit(clean(child.tail))

    def para(self, el, style=None, space_after=6, **kw):
        p = self.doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        self.runs(p, el, **kw)
        return p

    def text_para(self, txt, bold=False, size=None, color=None, space_after=6, italic=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(txt)
        r.bold, r.italic = bold, italic
        if size:
            r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        return p

    # ------------------------------------------------------------- bloques
    def table(self, el):
        rows = el.xpath('./tr | ./tbody/tr | ./thead/tr')
        if not rows:
            return
        ncols = max(len(r.xpath('./td | ./th')) for r in rows)
        # Las tablas con `border:0` son maquetación (texto junto a una imagen),
        # no datos: con rejilla parecerían un cuadro de contenido que no son.
        layout = 'border:0' in (el.get('style') or '')
        t = self.doc.add_table(rows=0, cols=ncols)
        if layout:
            no_borders(t)
        else:
            t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, tr in enumerate(rows):
            cells = tr.xpath('./td | ./th')
            row = t.add_row()
            header = cells[0].tag == 'th'
            for j in range(ncols):
                cell = row.cells[j]
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                if j < len(cells):
                    self.runs(cell.paragraphs[0], cells[j],
                              bold=header, color=PRIMARY if header else None,
                              size=9.5)
                if layout:
                    continue
                if header:
                    shade(cell, FILL_HEAD)
                elif i % 2 == 0:
                    shade(cell, FILL_ZEBRA)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def callout(self, el):
        variant = next((c for c in classes(el) if c in ('warn', 'ok', 'violet')), '')
        t = self.doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        shade(cell, CALLOUT_FILL[variant])
        no_borders(t)
        first = True
        for child in el:
            if 'lbl' in classes(child):
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                self.runs(p, child, bold=True, color=CALLOUT_INK[variant], size=9)
                first = False
            else:
                p = cell.paragraphs[0] if first else cell.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                self.runs(p, child, size=10)
                first = False
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def uc(self, el):
        head = el.find('./div[@class="uc-head"]')
        if head is None:
            head = next((c for c in el if 'uc-head' in classes(c)), None)
        if head is not None:
            violet = 'violet' in classes(head)
            t = self.doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            shade(cell, '6D3FC4' if violet else '00A39E')
            no_borders(t)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for div in head:
                if 'id' in classes(div):
                    self.runs(p, div, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9)
                elif 'title' in classes(div):
                    p2 = cell.add_paragraph()
                    p2.paragraph_format.space_after = Pt(0)
                    self.runs(p2, div, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=13)
            self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        body = next((c for c in el if 'uc-body' in classes(c)), None)
        if body is not None:
            self.walk(body)

    def uc_grid(self, el):
        pairs = [(d.find('./div[@class="k"]'), d.find('./div[@class="v"]')) for d in el]
        pairs = [(k, v) for k, v in pairs if k is not None and v is not None]
        if not pairs:
            return
        t = self.doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'
        for k, v in pairs:
            row = t.add_row()
            for cell, el2, bold, color in ((row.cells[0], k, True, MUTED),
                                           (row.cells[1], v, False, None)):
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                self.runs(p, el2, bold=bold, color=color, size=9)
            shade(row.cells[0], FILL_ZEBRA)
        t.columns[0].width = Cm(4.2)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def shots(self, el):
        for fig in el:
            img = fig.find('.//img')
            cap = fig.find('.//figcaption')
            if img is not None:
                src = os.path.join(DOCS, img.get('src'))
                if os.path.exists(src):
                    # Las capturas son de móvil (390×844): a 6 cm de ancho caben
                    # dos por página sin que la letra de la app deje de leerse.
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run().add_picture(src, width=Cm(6.4))
            if cap is not None:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(8)
                self.runs(p, cap, size=8.5, color=INK2)

    def flowmap(self, el):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        parts = [clean(n.text_content()).strip() for n in el if 'node' in classes(n)]
        r = p.add_run('  →  '.join(parts))
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = PRIMARY

    def toc(self, el):
        for li in el:
            if 'section-label' in classes(li):
                self.text_para(clean(li.text_content()).strip().upper(),
                               bold=True, size=9, color=PRIMARY, space_after=2)
                continue
            num = li.find('./span[@class="num"]')
            txt = li.find('./span[@class="txt"]')
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.6)
            if num is not None:
                r = p.add_run(clean(num.text_content()).strip().ljust(8))
                r.bold = True
                r.font.color.rgb = PRIMARY
            if txt is not None:
                p.add_run(clean(txt.text_content()).strip())

    def cover(self, el):
        def first(cls):
            found = el.xpath('.//*[contains(concat(" ", @class, " "), " %s ")]' % cls)
            return found[0] if found else None

        self.text_para('valeria+', bold=True, size=20, color=PRIMARY, space_after=14)
        mascota = os.path.join(DOCS, 'lua-mascota.png')
        if os.path.exists(mascota):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
            p.add_run().add_picture(mascota, width=Cm(4.0))
        for cls, size, bold, color in (('pill', 10, True, PRIMARY),
                                       ('sub', 11, False, INK2),
                                       ('meta', 9.5, False, MUTED)):
            node = first(cls)
            if node is not None:
                self.para(node, size=size, bold=bold, color=color, space_after=10)
        h1 = el.find('.//h1')
        if h1 is not None:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            # El <br> del título no deja espacio en text_content(): sin esto
            # sale "Manual deCasos de Uso".
            r = p.add_run(clean(' '.join(h1.itertext())).strip())
            r.bold = True
            r.font.size = Pt(30)
            r.font.color.rgb = INK
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------------------------------------------------------- walk
    def walk(self, parent):
        for el in parent:
            if not isinstance(el.tag, str):
                continue  # comentarios
            cls = classes(el)
            tag = el.tag

            if 'cover' in cls:
                self.cover(el)
            elif 'chapter-kicker' in cls:
                self.text_para(clean(el.text_content()).strip().upper(),
                               bold=True, size=8.5, color=PRIMARY, space_after=1)
            elif tag == 'h2':
                self.doc.add_heading(clean(el.text_content()).strip(), level=1)
            elif tag == 'h3':
                self.doc.add_heading(clean(el.text_content()).strip(), level=2)
            elif tag == 'h4':
                self.doc.add_heading(clean(el.text_content()).strip(), level=3)
            elif 'toc' in cls:
                self.toc(el)
            elif 'flowmap' in cls:
                self.flowmap(el)
            elif 'callout' in cls:
                self.callout(el)
            elif 'uc' in cls:
                self.uc(el)
            elif 'uc-grid' in cls:
                self.uc_grid(el)
            elif 'shots' in cls:
                self.shots(el)
            elif 'footer-note' in cls:
                self.para(el, size=8.5, color=MUTED, italic=True, space_after=2)
            elif tag == 'table':
                self.table(el)
            elif tag == 'p':
                small = 'small' in cls
                self.para(el, size=9 if small else None,
                          color=MUTED if 'muted' in cls else None)
            elif tag in ('ul', 'ol'):
                style = 'List Number' if tag == 'ol' else 'List Bullet'
                for li in el:
                    p = self.doc.add_paragraph(style=style)
                    p.paragraph_format.space_after = Pt(2)
                    self.runs(p, li)
            elif tag == 'section':
                if 'page-break' in cls:
                    self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                self.walk(el)
            elif tag in ('div', 'figure'):
                self.walk(el)


def build():
    doc = Document()
    doc.core_properties.title = 'Valeria+ · Manual de Casos de Uso'
    doc.core_properties.author = 'Proyecto Valeria+'
    doc.core_properties.language = 'es-ES'

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(2.2)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    for name, size, color, before in (('Heading 1', 17, PRIMARY, 16),
                                      ('Heading 2', 13, INK, 12),
                                      ('Heading 3', 11.5, INK2, 10)):
        st = doc.styles[name]
        st.font.name = 'Calibri'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(5)

    html = lxml.html.parse(SRC).getroot()
    b = Builder(doc)
    b.walk(html.find('body'))
    doc.save(OUT)
    print('escrito %s · %d figuras' % (os.path.relpath(OUT, ROOT), b.figura))


if __name__ == '__main__':
    build()
